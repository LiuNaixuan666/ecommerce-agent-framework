from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


OUT_DIR = Path("output/dissertation")
OUT_DIR.mkdir(parents=True, exist_ok=True)
OUT_PATH = OUT_DIR / "Liu_Naixuan_Dissertation_Initial_Draft_2026_CN.docx"


def set_font(run, size=12, bold=None, italic=None):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def p(doc, text="", style=None, align=None, bold=False, size=12):
    para = doc.add_paragraph(style=style)
    if align is not None:
        para.alignment = align
    run = para.add_run(text)
    set_font(run, size=size, bold=bold)
    return para


def h(doc, text, level=1):
    para = doc.add_heading(text, level=level)
    for run in para.runs:
        set_font(run, size=16 if level == 1 else 14 if level == 2 else 12, bold=True)
    return para


def bullets(doc, items):
    for item in items:
        p(doc, item, style="List Bullet")


def numbered(doc, items):
    for item in items:
        p(doc, item, style="List Number")


def table(doc, headers, rows):
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.style = "Table Grid"
    for i, header in enumerate(headers):
        tbl.rows[0].cells[i].text = header
    for row in rows:
        cells = tbl.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
    doc.add_paragraph("")
    for row in tbl.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                para.paragraph_format.line_spacing = 1.15
                para.paragraph_format.space_after = Pt(2)
                for run in para.runs:
                    set_font(run, size=10)
    return tbl


def build():
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(1)
    sec.bottom_margin = Inches(1)
    sec.left_margin = Inches(1)
    sec.right_margin = Inches(1)

    doc.styles["Normal"].font.name = "Times New Roman"
    doc.styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    doc.styles["Normal"].font.size = Pt(12)
    doc.styles["Normal"].paragraph_format.line_spacing = 1.5
    doc.styles["Normal"].paragraph_format.space_after = Pt(6)

    footer = sec.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = footer.add_run("DASE7099 毕业论文初稿 - LIU Naixuan")
    set_font(r, 9)

    p(doc, "香港大学", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
    p(doc, "数据与系统工程系", align=WD_ALIGN_PARAGRAPH.CENTER)
    p(doc, "工业工程与物流管理理学硕士", align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph("")
    p(doc, "面向可扩展电商客服的本地多平台 Agentic RAG 框架", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=18)
    p(doc, "以拼多多页面自动化与商品知识库问答为例", align=WD_ALIGN_PARAGRAPH.CENTER, size=14)
    doc.add_paragraph("")
    p(doc, "毕业论文初稿", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
    doc.add_paragraph("")
    p(doc, "LIU Naixuan（3036565551）", align=WD_ALIGN_PARAGRAPH.CENTER)
    p(doc, "导师：Dr. S.H. CHOI", align=WD_ALIGN_PARAGRAPH.CENTER)
    p(doc, "提交日期：2026年6月18日", align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_page_break()

    h(doc, "摘要")
    for text in [
        "随着电商平台和直播电商的发展，商家每天需要处理大量重复性客服问题，包括商品介绍、库存状态、发货时效、运费规则、退换货政策、订单状态和售后风险等。现有平台自带的智能客服工具通常只能服务单个平台，商家如果同时经营拼多多、闲鱼、淘宝、京东、抖音等渠道，就需要在不同系统中重复配置商品知识、回复策略和售后规则，使用成本高，维护效率低，并且容易出现不同平台回复口径不一致的问题。",
        "本论文提出并实现一个本地多平台 Agentic RAG 电商客服框架。系统不以直接申请平台官方 API 作为第一路径，而是采用本地浏览器自动化方式读取电商平台网页版客服窗口中的买家消息和商品上下文，再调用本地 AI 服务生成结构化回复。后端基于 FastAPI，知识库采用本地向量检索，工作流包含意图识别、知识检索、回复生成、质量闸门、风险判断和自动发送边界。系统返回的不只是客服话术，还包括风险等级、证据来源、阻断原因、是否允许自动发送、是否需要转人工等字段，使自动化回复可以被显式控制。",
        "项目最初以图书品类 RAG 智能客服为研究对象，后续逐步演进为面向真实商家工作流的本地多平台客服中台。目前系统已经完成后端稳定化、PDF/DOCX/CSV 等文档上传解析、结构化客服回复、RPA 标准接口、本地 Local Agent 骨架、Mock 客服闭环、浏览器页面适配原型以及拼多多页面 dry-run 验证。已验证的拼多多案例中，Local Agent 能够读取真实客服页面中的最新买家消息“你好”，后端识别为低风险问候类消息并生成“您好，请问有什么可以帮您？”的回复，同时由于未开启真实发送权限，系统按安全策略跳过发送。",
        "本研究的贡献在于提出一种不依赖官方平台 API 的本地多平台 AI 客服实现路径，并通过 Agentic RAG 和风险控制机制解决商品知识绑定、回复可信度、自动发送安全和人工接管等问题。后续工作将继续完善拼多多真实页面适配、商家数据库接入、商品知识精细绑定、问答模板学习以及多平台扩展，并通过标准测试集评估回答准确率、检索相关性、响应延迟和转人工策略效果。",
    ]:
        p(doc, text)
    doc.add_page_break()

    h(doc, "声明")
    p(doc, "本人 LIU Naixuan 声明，本毕业论文初稿《面向可扩展电商客服的本地多平台 Agentic RAG 框架》为本人在导师指导下完成的工作，除文中已注明引用的内容外，未曾以相同或类似形式提交至其他机构用于获得学位、文凭或其他资格。")
    p(doc, "签名：__________________________")
    p(doc, "日期：2026年6月18日")
    doc.add_page_break()

    h(doc, "致谢")
    p(doc, "感谢 Dr. S.H. CHOI 在本项目选题、系统设计和论文推进过程中给予的指导与建议。感谢香港大学数据与系统工程系提供的学习环境和项目支持。也感谢在系统原型测试和需求讨论中给予反馈的同学、朋友和潜在用户。")
    doc.add_page_break()

    h(doc, "目录")
    bullets(doc, ["摘要", "声明", "致谢", "第1章 绪论", "第2章 文献综述", "第3章 研究方法", "第4章 系统设计与实现", "第5章 评估与初步结果", "第6章 讨论", "第7章 结论与未来工作", "参考文献", "附录"])
    doc.add_page_break()

    chapter_1(doc)
    chapter_2(doc)
    chapter_3(doc)
    chapter_4(doc)
    chapter_5(doc)
    chapter_6(doc)
    chapter_7(doc)

    for para in doc.paragraphs:
        if not para.style.name.startswith("Heading") and para.text.strip():
            para.paragraph_format.line_spacing = 1.5
            para.paragraph_format.space_after = Pt(6)
            if not para.style.name.startswith("List"):
                para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        for run in para.runs:
            if run.font.size is None:
                set_font(run)

    doc.save(OUT_PATH)
    print(OUT_PATH)


def chapter_1(doc):
    h(doc, "第1章 绪论")
    h(doc, "1.1 研究背景", 2)
    p(doc, "电商客服是连接商家、平台和消费者的重要环节。对中小商家而言，客服问题高度重复，但又不能完全用固定模板解决。例如，客户可能询问商品是否有货、尺寸如何选择、是否包邮、多久发货、能否退换、订单为什么没有更新，或者对售后处理表达不满。传统关键词匹配或固定话术系统难以处理多样表达，而通用大模型虽然语言能力强，却容易在缺少商家知识时编造事实。")
    p(doc, "因此，本项目采用 RAG 与 Agent 工作流结合的方式，将商家上传的商品文档、政策文档、问答模板和未来接入的数据库作为知识来源，让大模型在有证据的范围内生成回复。同时，由于平台官方客服 API 获取门槛较高，本研究将本地浏览器自动化作为第一阶段的平台接入路径，即通过本地 Local Agent 读取网页版客服窗口信息，再调用本地 AI 服务完成回复生成与风险判断。")
    h(doc, "1.2 问题定义", 2)
    bullets(doc, [
        "多平台配置重复：商家在不同电商平台需要重复维护商品信息、客服规则和回复模板。",
        "官方 API 权限受限：消息、订单和自动发送接口通常需要平台审核或服务商资质。",
        "回答质量不稳定：单纯依赖大模型容易产生不符合商品事实或店铺政策的回复。",
        "自动发送存在风险：退款、投诉、差评、赔偿等高风险问题不应由 AI 无条件自动处理。",
        "网页适配不稳定：RPA 或浏览器自动化依赖页面结构，需要可配置、可观测、可回退的适配机制。",
    ])
    h(doc, "1.3 研究目标", 2)
    numbered(doc, [
        "构建一个本地运行的多平台电商客服框架，使商家可以统一维护知识库并复用到不同平台。",
        "实现意图驱动的 Agentic RAG 工作流，支持商品咨询、政策咨询、订单服务、闲聊、模糊问题和高风险售后识别。",
        "设计结构化客服回复格式，使 Local Agent 能根据 auto_send_allowed、risk_level 和 blockers 决定是否自动发送或转人工。",
        "以拼多多为第一个真实平台适配目标，验证读取买家消息、生成回复、dry-run 安全跳过发送的完整链路。",
        "为后续接入商家数据库、问答模板学习、多平台控制台和真实自动发送验证奠定基础。",
    ])
    h(doc, "1.4 研究贡献", 2)
    p(doc, "本研究的主要贡献包括：提出无官方 API 情况下的本地多平台客服中台架构；将 RAG、意图识别、风险策略和浏览器自动化结合到一个可运行原型中；将客服回复从单一文本扩展为结构化决策结果；并通过拼多多 dry-run 验证证明该路线具备实际可行性。")


def chapter_2(doc):
    h(doc, "第2章 文献综述")
    h(doc, "2.1 电商智能客服", 2)
    p(doc, "早期智能客服多依赖 FAQ、关键词匹配和规则树。这类方法可控性强，但面对复杂表达和上下文变化时灵活性不足。随着大模型发展，客服系统可以更自然地理解用户问题并生成回复，但电商场景对事实准确性、政策一致性和风险控制要求更高，因此不能只依赖通用模型。")
    h(doc, "2.2 RAG 与知识增强生成", 2)
    p(doc, "RAG 的核心思想是在生成答案前先检索外部知识，使模型回答基于可追溯证据。对电商客服而言，RAG 可以把商品说明、退货政策、运费规则、FAQ 和库存表纳入回答依据。但 RAG 也存在检索错误、商品绑定不清、文档粒度不合理等问题，因此需要配合意图识别和元数据管理。")
    h(doc, "2.3 Agentic Workflow", 2)
    p(doc, "Agentic Workflow 并不意味着让模型完全自主行动，而是将任务拆分为可控步骤。本项目中的 Agent 先识别问题意图，再选择检索范围，随后生成回复，并通过质量闸门和风险策略决定是否允许自动发送。这种结构比单一 prompt 更容易维护和扩展，也更适合带有业务风险的客服场景。")
    h(doc, "2.4 无 API 平台适配", 2)
    p(doc, "理想情况下，平台官方 API 可以稳定提供消息、订单和商品数据。但现实中，很多平台 API 需要服务商资质或商家授权。对于毕业设计和早期产品验证，浏览器自动化是一种更现实的路线。它通过模拟人工客服正在使用的网页来读取消息、填写回复和点击发送，但也需要处理页面变化、电脑在线、登录状态和自动发送安全等问题。")


def chapter_3(doc):
    h(doc, "第3章 研究方法")
    h(doc, "3.1 设计科学方法", 2)
    p(doc, "本论文采用设计科学研究方法，目标是构建一个能解决实际问题的软件原型，而不仅是理论分析。研究过程包括需求识别、系统架构设计、模块实现、功能验证和迭代优化。项目从最初的图书品类 RAG 客服，逐步演进为本地多平台客服框架。")
    h(doc, "3.2 系统需求", 2)
    table(doc, ["需求", "说明", "对应模块"], [
        ["统一知识库", "商家一次上传商品、政策和问答模板，多平台复用", "文档解析、向量库、元数据"],
        ["无 API 接入", "不依赖官方客服 API，先通过浏览器页面读取消息", "Local Agent、平台 Adapter"],
        ["结构化回复", "输出话术、风险、证据、阻断原因和自动发送决策", "RPA schema、workflow"],
        ["自动发送边界", "只有低风险、证据充分、消息新鲜且显式允许时才发送", "risk policy、quality gate"],
        ["转人工机制", "高风险、无答案、模糊问题进入人工处理", "handoff_required、blockers"],
    ])
    h(doc, "3.3 Agentic RAG 流程", 2)
    numbered(doc, [
        "接收买家消息、平台、会话、商品上下文和订单上下文。",
        "通过意图解析器判断用户是在咨询商品、政策、订单，还是提出投诉或模糊问题。",
        "根据意图选择商品知识、政策文档、FAQ 模板或数据库作为检索来源。",
        "生成客服回复，并保留证据来源和置信度信息。",
        "通过质量闸门判断是否存在弱检索、无证据、答非所问或需要澄清的情况。",
        "通过风险策略判断是否允许自动发送，或必须转人工。",
        "返回结构化结果给 RPA 或 Local Agent。",
    ])
    h(doc, "3.4 评估方法", 2)
    p(doc, "评估分为三个层次：第一是后端模块测试，包括健康检查、上传、解析、检索、意图识别和风险策略；第二是 Mock 客服闭环测试，验证 Local Agent 能读取消息、调用后端和记录发送结果；第三是真实平台 dry-run 验证，先在不点击发送的情况下确认拼多多页面消息读取和 AI 回复链路。最终评估将补充标准问题集，对回答准确性、检索相关性、自动发送决策和响应延迟进行量化。")


def chapter_4(doc):
    h(doc, "第4章 系统设计与实现")
    h(doc, "4.1 总体架构", 2)
    p(doc, "系统由四层组成：前端管理层、Agent 工作流层、知识库层和执行层。前端用于上传知识、测试对话、查看平台状态和配置 AI 客服模式；Agent 工作流负责意图识别、检索、生成和风险控制；知识库层存储文档切块和向量索引；执行层通过 Local Agent 和平台 Adapter 与浏览器页面交互。")
    h(doc, "4.2 后端服务", 2)
    p(doc, "后端基于 FastAPI，提供 /health、/api/chat/health、/api/knowledge/health 等健康检查接口，也提供 /api/chat/rpa/message 和 /api/chat/rpa/send-result 供 RPA 或 Local Agent 调用。Local Agent 还通过 heartbeat 和 status 接口向系统报告当前平台连接状态。")
    h(doc, "4.3 知识库与文档上传", 2)
    p(doc, "知识库支持 PDF、DOCX、CSV 和文本类文件解析。商品文档、政策文档和 FAQ 模板被切分后写入 ChromaDB 向量库。后续将进一步强化元数据，包括 platform、product_id、sku_id、doc_type 和 scope，确保当前商品的问题优先检索对应商品的知识，而不是误用其他商品或全店通用政策。")
    h(doc, "4.4 风险策略与自动发送", 2)
    p(doc, "系统将自动发送视为高风险动作，因此默认 dry-run，不真实点击发送按钮。只有当消息为最新买家消息、风险等级低、证据充分、质量闸门通过，并且启动命令显式打开 allow-real-send 时，Local Agent 才允许执行真实发送。退款、退货、投诉、赔偿、差评、假货和平台违规等问题默认进入人工处理。")
    h(doc, "4.5 Local Agent 与拼多多 Adapter", 2)
    p(doc, "Local Agent 负责持续监听平台页面，读取最新买家消息，收集商品上下文，调用本地后端，并根据结构化结果决定 dry-run、发送、失败重试或转人工。目前已实现 BasePlatformAdapter、MockShopAdapter、GenericWebChatAdapter 和拼多多本地选择器配置 pinduoduo_web.local.json。拼多多页面验证已经跑通读取真实消息并生成回复的 dry-run 链路。")
    h(doc, "4.6 前端产品方向", 2)
    p(doc, "前端将从单一聊天页面升级为多平台控制台。主页显示拼多多、闲鱼、淘宝/千牛、京东、抖音等平台入口，点击平台后进入对应工作台，展示登录状态、Local Agent 状态、买家消息、商品上下文、AI 回复、风险等级和处理结果。同时，前端需要支持知识库管理、AI 模式选择、转人工规则配置和问答模板优化。")


def chapter_5(doc):
    h(doc, "第5章 评估与初步结果")
    h(doc, "5.1 当前验证状态", 2)
    p(doc, "当前阶段的评估重点是功能可行性，而不是最终大规模实验结果。系统已经完成后端启动与健康检查、文档上传解析、RAG 问答链路、Mock 客服闭环、Local Agent 监听流程、风险策略和拼多多 dry-run 验证。")
    h(doc, "5.2 拼多多 dry-run 案例", 2)
    p(doc, "在真实拼多多商家客服页面中，Local Agent 读取到最新买家消息“你好”，后端将其识别为问候类低风险消息，并生成回复“您好，请问有什么可以帮您？”。由于当时 dry-run=true 且 allow-real-send=false，系统没有点击发送按钮，而是记录为 skipped_dry_run。该案例验证了无官方 API 情况下“读取页面消息 -> 调用本地 AI -> 生成结构化回复 -> 安全跳过真实发送”的完整路径。")
    h(doc, "5.3 风险策略验证", 2)
    table(doc, ["问题类型", "期望行为", "当前状态"], [
        ["问候语", "低风险，可生成简短回复；dry-run 下不发送", "已在拼多多页面验证"],
        ["商品功能", "优先引用商品知识库回答", "后端链路已验证，需扩展样本"],
        ["运费/退货政策", "必须基于政策文档回答，否则澄清或转人工", "待最终测试集验证"],
        ["退款/退货/换货", "中风险，默认需要更强证据或人工介入", "规则已实现"],
        ["投诉/赔偿/差评/假货", "高风险，不允许自动发送", "规则已实现"],
        ["知识库无答案", "不得编造，应拒答、澄清或转人工", "质量闸门已实现，需量化评估"],
    ])
    h(doc, "5.4 后续量化评估计划", 2)
    p(doc, "最终论文将构建包含商品咨询、政策咨询、订单服务、模糊问题、无答案问题和高风险售后的测试集。每个样本标注意图、标准答案、证据来源和期望自动发送决策。系统将与 LLM-only 和普通 RAG 基线比较，指标包括意图准确率、回答正确率、证据命中率、自动发送决策准确率、澄清质量和响应延迟。")


def chapter_6(doc):
    h(doc, "第6章 讨论")
    h(doc, "6.1 路线演进", 2)
    p(doc, "项目从最初的图书品类 RAG 客服，演进为面向真实平台工作流的本地多平台客服中台。这一变化来自实践中的关键发现：电商客服的难点不仅是模型回答质量，更包括平台接入、商品上下文、发送权限、风险控制和商家配置效率。")
    h(doc, "6.2 本地方案优势", 2)
    bullets(doc, [
        "不依赖平台官方 API，适合毕业设计和早期产品验证。",
        "商家可以统一维护知识库，减少多平台重复配置。",
        "支持全托管、半托管和智能转接等不同 AI 客服模式。",
        "通过 dry-run、最新消息过滤、风险闸门和显式发送权限降低误发风险。",
        "平台 Adapter 可插拔，后续可逐步扩展到闲鱼、淘宝、京东、抖音等平台。",
    ])
    h(doc, "6.3 局限性", 2)
    p(doc, "本地浏览器自动化依赖电脑和浏览器保持在线，页面必须处于可访问状态；如果平台改版，选择器可能失效；如果订单或库存信息没有出现在页面上，系统需要额外接入商家数据库。与官方 API 相比，本地自动化更适合验证和轻量部署，但稳定性和规模化能力仍需进一步提升。")
    h(doc, "6.4 知识库持续优化", 2)
    p(doc, "未来系统应支持从人工客服中学习。当 AI 不知道如何回答而转人工后，商家确认过的人工回复可以进入问答模板库。这样系统可以不断积累真实场景话术，减少重复转人工次数，同时保证新知识经过商家确认。")


def chapter_7(doc):
    h(doc, "第7章 结论与未来工作")
    h(doc, "7.1 结论", 2)
    p(doc, "本文提出并实现了一个本地多平台 Agentic RAG 电商客服框架，用于解决商家在多平台经营中智能客服配置重复、官方 API 获取困难、回复口径不一致和自动发送风险高等问题。系统通过本地知识库、意图驱动检索、结构化回复、风险策略和 Local Agent 浏览器自动化，验证了一条不依赖官方 API 的可行实现路径。")
    p(doc, "当前原型已经完成从后端稳定化到拼多多页面 dry-run 的核心闭环。虽然真实自动发送、多客户并发处理、商家数据库接入和多平台适配仍需进一步完善，但系统已经证明：只要将 AI 回复和执行动作分离，并用明确的风险字段控制自动化边界，本地 AI 客服中台可以逐步接入真实电商工作流。")
    h(doc, "7.2 未来工作", 2)
    bullets(doc, [
        "完善多平台控制台和拼多多平台详情页。",
        "补充拼多多真实页面的多轮、多客户、商品卡片和售后场景测试。",
        "实现商家数据库可视化接入，用于库存、订单和商品信息查询。",
        "加强商品文档与 product_id/SKU 的绑定。",
        "实现人工确认回复转问答模板的学习闭环。",
        "扩展闲鱼、淘宝/千牛、京东和抖音等平台 Adapter。",
        "完成最终量化实验并补充论文图表和统计结果。",
    ])
    h(doc, "参考文献")
    for ref in [
        "Lewis, P. et al. (2020). Retrieval-Augmented Generation for Knowledge-Intensive NLP Tasks.",
        "OpenAI. (2023). GPT-4 Technical Report.",
        "Qiu, M. et al. (2017). AliMe Chat: A Sequence to Sequence and Rerank Based Chatbot Engine.",
        "Radziwill, N., & Benton, M. (2017). Evaluating Quality of Chatbots and Intelligent Conversational Agents.",
        "Zhou, L. et al. (2020). The Design and Implementation of XiaoIce, an Empathetic Social Chatbot.",
    ]:
        p(doc, ref)
    h(doc, "附录")
    bullets(doc, [
        "主要接口：/health、/api/chat/health、/api/knowledge/health、/api/chat/rpa/message、/api/chat/rpa/send-result。",
        "当前阶段：系统稳定化、结构化回复、风险策略、Local Agent、Mock 闭环和拼多多 dry-run 已完成。",
        "待更新内容：最终实验数据、系统截图、架构图、目录页码和参考文献格式。",
    ])


if __name__ == "__main__":
    build()
